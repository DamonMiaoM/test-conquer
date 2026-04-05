#!/usr/bin/env python3
"""从Word文档提取单选题和多选题，生成格式化的Markdown文件"""
import re
from docx import Document

DOCX_PATH = "/Users/Damon/Projects/Test Conquer/docs/第3部分3级_理论知识复习题 2_w.docx"
SINGLE_OUT = "/Users/Damon/Projects/Test Conquer/docs/单选题.md"
MULTI_OUT = "/Users/Damon/Projects/Test Conquer/docs/多选题.md"

# Read answer file for reference
ANSWER_PATH = "/Users/Damon/Projects/Test Conquer/docs/人工智能训练师（三级）技能操作复习答案汇总_lint.md"
with open(ANSWER_PATH, 'r', encoding='utf-8') as f:
    answer_text = f.read()

def extract_answers_for_type(section_text, start_marker, q_type='single'):
    """Extract answers from the answer reference file"""
    answers = {}
    # Find the section
    start_idx = section_text.find(start_marker)
    if start_idx == -1:
        return answers
    
    # For single choice: pattern like "1. ...C)...", "2. ...A)..."
    # For multi choice: pattern like "1. ...ABCD)...", "职业道德...ABC..."
    
    return answers

# Parse single choice answers from reference
single_answers = {}
# Pattern: number. ... (X) ... or number\. ...\s+\(X\)
for m in re.finditer(r'(?:(?:^|\n)\s*)(\d+)[\.。、]\s*.*?\(([A-D])\)', answer_text):
    num = int(m.group(1))
    ans = m.group(2)
    single_answers[num] = ans

# Also try pattern with Chinese period
for m in re.finditer(r'(\d+)[。、]\s*.*?\(([A-D])\)', answer_text):
    num = int(m.group(1))
    ans = m.group(2)
    if num not in single_answers:
        single_answers[num] = ans

# Parse multi choice answers from reference
multi_answers = {}
# Find multi-choice section
multi_section_start = answer_text.find('多选题')
if multi_section_start > 0:
    multi_text = answer_text[multi_section_start:]
    # Pattern for multi: can be ABCDE, ABCD, etc.
    for m in re.finditer(r'(\d+)[\.。、]\s*.*?\(([A-E]+)\)', multi_text):
        num = int(m.group(1))
        ans = m.group(2)
        multi_answers[num] = ans
    # Also try without number prefix for the first few questions
    # Pattern: text followed by (ABCDE) etc at end of line
    for m in re.finditer(r'([A-Z]+)\)\s*$', multi_text, re.MULTILINE):
        ans_str = m.group(1)
        if len(ans_str) >= 2 and ans_str.isalpha():
            pass  # These are harder to associate without numbers

doc = Document(DOCX_PATH)

# Find section boundaries
single_start = None
multi_start = None
judge_start = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '单选题' in text and '选择一个正确的答案' in text:
        single_start = i + 1
    elif '多选题' in text and '选择多个正确的答案' in text:
        multi_start = i + 1
    elif '判断题' in text and '判断结果' in text:
        judge_start = i + 1

print(f"Single choice: paras {single_start}-{multi_start-1}")
print(f"Multi choice: paras {multi_start}-{judge_start-1 if judge_start else 'end'}")

def is_option_line(text):
    """Check if line is an option like (A)xxx"""
    return bool(re.match(r'^\s*\([A-E]\)', text))

def get_option_letter(text):
    """Get option letter from line"""
    m = re.match(r'^\s*\(([A-E])\)', text)
    return m.group(1) if m else None

def is_question_start(text, q_type='single'):
    """Check if line starts a new question"""
    # Pattern: "number. text" or "number text"
    if q_type == 'multi':
        # Multi choice may not have numbers for first few
        # Check if it's a stem line (not option, not empty)
        if is_option_line(text):
            return False
        if not text.strip():
            return False
        # If starts with a number, definitely a question
        if re.match(r'^\s*\d+[\.\。、]', text):
            return True
        # Otherwise, it's a question stem if it contains Chinese or parentheses
        if re.search(r'[\u4e00-\u9fff]', text) and not is_option_line(text):
            return True
        return False
    else:
        return bool(re.match(r'^\s*\d+[\.\。、]', text))

def parse_questions(start_para, end_para, q_type='single'):
    """Parse questions from docx paragraphs"""
    questions = []
    current_num = 0
    current_stem = ''
    current_options = []
    
    for i in range(start_para, end_para):
        text = doc.paragraphs[i].text.strip()
        if not text:
            continue
        
        if is_option_line(text):
            # This is an option line
            letter = get_option_letter(text)
            content = re.sub(r'^\s*\([A-E]\)\s*', '', text)
            current_options.append((letter, content))
        elif is_question_start(text, q_type):
            # Save previous question
            if current_stem or current_options:
                questions.append({
                    'num': current_num,
                    'stem': current_stem,
                    'options': current_options
                })
            
            # Parse new question
            m = re.match(r'^\s*(\d+)[\.\。、]\s*(.*)', text)
            if m:
                current_num = int(m.group(1))
                current_stem = m.group(2)
            else:
                current_num += 1
                current_stem = text
            current_options = []
        else:
            # Continuation of stem or orphan line
            if current_options:
                # Options already started, this might be orphan text
                # or continuation. Skip for now.
                pass
            else:
                current_stem += ' ' + text
    
    # Don't forget last question
    if current_stem or current_options:
        questions.append({
            'num': current_num,
            'stem': current_stem,
            'options': current_options
        })
    
    return questions

# Parse single choice questions
single_qs = parse_questions(single_start, multi_start, 'single')
print(f"\nParsed {len(single_qs)} single choice questions")

# Check for missing options
missing_opts = 0
for q in single_qs:
    if len(q['options']) < 4:
        missing_opts += 1
        print(f"  Q{q['num']}: only {len(q['options'])} options: {q['stem'][:60]}")
print(f"Total with missing options: {missing_opts}")

# Parse multi choice questions
multi_end = len(doc.paragraphs)
multi_qs = parse_questions(multi_start, multi_end, 'multi')
print(f"\nParsed {len(multi_qs)} multi choice questions")

# Check for missing options in multi
missing_opts_m = 0
for q in multi_qs:
    if len(q['options']) < 3:
        missing_opts_m += 1
        print(f"  Q{q['num']}: only {len(q['options'])} options: {q['stem'][:60]}")
print(f"Total with missing options: {missing_opts_m}")

# Generate single choice markdown
def generate_single_md(questions, answers):
    lines = []
    lines.append('# 人工智能训练师（三级）')
    lines.append('')
    lines.append('## 单选题')
    lines.append('')
    lines.append('（选择一个正确的答案，将相应的字母填入题内的括号中。）')
    lines.append('')
    
    for idx, q in enumerate(questions, 1):
        # Clean stem: remove answer placeholders
        stem = q['stem']
        # Replace (    ) or () with answer
        answer = answers.get(q['num'], '')
        
        # Check if stem already has answer
        stem_has_answer = bool(re.search(r'\(\s*[A-D]\s*\)', stem))
        
        if stem_has_answer:
            # Keep existing answer
            pass
        elif answer:
            # Insert answer
            stem = re.sub(r'\(\s*\)', f'({answer})', stem)
            if not re.search(r'\([A-D]\)', stem):
                # Append answer if no placeholder found
                stem = stem.rstrip(' .') + f'({answer})'
        
        # Renumber
        opt_str = ' '.join(f'({letter}){content}' for letter, content in q['options'])
        
        lines.append(f'{idx}. {stem}')
        lines.append('')
        if opt_str:
            lines.append(f'   {opt_str}')
            lines.append('')
    
    return '\n'.join(lines)

def generate_multi_md(questions, answers):
    lines = []
    lines.append('# 人工智能训练师（三级）')
    lines.append('')
    lines.append('## 多选题')
    lines.append('')
    lines.append('（选择多个正确的答案，将相应的字母填入题内的括号中。）')
    lines.append('')
    
    for idx, q in enumerate(questions, 1):
        stem = q['stem']
        answer = answers.get(q['num'], answers.get(idx, ''))
        
        stem_has_answer = bool(re.search(r'\(\s*[A-E]+\s*\)', stem))
        
        if stem_has_answer:
            pass
        elif answer:
            stem = re.sub(r'\(\s*\)', f'({answer})', stem)
            if not re.search(r'\([A-E]+\)', stem):
                stem = stem.rstrip(' .') + f'({answer})'
        
        opt_str = ' '.join(f'({letter}){content}' for letter, content in q['options'])
        
        if q['num'] > 0:
            lines.append(f'{idx}. {stem}')
        else:
            lines.append(f'{idx}. {stem}')
        lines.append('')
        if opt_str:
            lines.append(f'   {opt_str}')
            lines.append('')
    
    return '\n'.join(lines)

# Generate and write
single_md = generate_single_md(single_qs, single_answers)
with open(SINGLE_OUT, 'w', encoding='utf-8') as f:
    f.write(single_md)
print(f"\nSingle choice written to {SINGLE_OUT}")

multi_md = generate_multi_md(multi_qs, multi_answers)
with open(MULTI_OUT, 'w', encoding='utf-8') as f:
    f.write(multi_md)
print(f"Multi choice written to {MULTI_OUT}")
